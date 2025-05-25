import requests
import sqlite3
import json
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime

# --- PARTE 1: Extração de Dados via API REST (Países) ---

def configurar_banco_dados_paises():
    conn = sqlite3.connect('paises.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS paises (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome_comum TEXT,
            nome_oficial TEXT,
            capital TEXT,
            continente TEXT,
            regiao TEXT,
            sub_regiao TEXT,
            populacao INTEGER,
            area REAL,
            nome_moeda TEXT,
            simbolo_moeda TEXT,
            idioma_principal TEXT,
            fuso_horario TEXT,
            url_bandeira TEXT
        )
    ''')
    conn.commit()
    conn.close()
    print("Banco de dados 'paises.db' e tabela 'paises' verificados/criados.")

def obter_dados_pais(nome_pais):
    url = f"https://restcountries.com/v3.1/name/{nome_pais}"
    try:
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()
        if not data:
            print(f"Nenhum dado encontrado para {nome_pais}.")
            return None

        pais = data[0]

        common_name = pais.get('name', {}).get('common', 'N/A')
        official_name = pais.get('name', {}).get('official', 'N/A')
        capital = pais.get('capital', ['N/A'])[0] if pais.get('capital') else 'N/A'
        continente = pais.get('continents', ['N/A'])[0] if pais.get('continents') else 'N/A'
        regiao = pais.get('region', 'N/A')
        sub_regiao = pais.get('subregion', 'N/A')
        populacao = pais.get('population', 0)
        area = pais.get('area', 0.0)

        moedas = pais.get('currencies', {})
        nome_moeda = 'N/A'
        simbolo_moeda = 'N/A'
        if moedas:
            codigo_moeda = list(moedas.keys())[0]
            nome_moeda = moedas[codigo_moeda].get('name', 'N/A')
            simbolo_moeda = moedas[codigo_moeda].get('symbol', 'N/A') 

        idiomas = pais.get('languages', {})
        idioma_principal = 'N/A'
        if idiomas:
            idioma_principal = list(idiomas.values())[0]

        fusos_horarios = pais.get('timezones', ['N/A'])
        fuso_horario = ', '.join(fusos_horarios) if fusos_horarios else 'N/A'
        
        flag_url = pais.get('flags', {}).get('png', 'N/A')

        return {
            "nome_comum": common_name,
            "nome_oficial": official_name,
            "capital": capital,
            "continente": continente,
            "regiao": regiao,
            "sub_regiao": sub_regiao,
            "populacao": populacao,
            "area": area,
            "nome_moeda": nome_moeda,
            "simbolo_moeda": simbolo_moeda,
            "idioma_principal": idioma_principal,
            "fuso_horario": fuso_horario,
            "url_bandeira": flag_url
        }

    except requests.exceptions.RequestException as e:
        print(f"Erro ao buscar dados para {nome_pais}: {e}")
        return None
    except json.JSONDecodeError:
        print(f"Erro ao decodificar JSON para {nome_pais}.")
        return None
    except IndexError:
        print(f"Nenhum dado de país válido encontrado na resposta para {nome_pais}.")
        return None

def inserir_dados_pais(dados):
    conn = sqlite3.connect('paises.db')
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO paises (
            nome_comum, nome_oficial, capital, continente, regiao, sub_regiao,
            populacao, area, nome_moeda, simbolo_moeda, idioma_principal,
            fuso_horario, url_bandeira
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        dados['nome_comum'], dados['nome_oficial'], dados['capital'],
        dados['continente'], dados['regiao'], dados['sub_regiao'],
        dados['populacao'], dados['area'], dados['nome_moeda'],
        dados['simbolo_moeda'], dados['idioma_principal'],
        dados['fuso_horario'], dados['url_bandeira']
    ))
    conn.commit()
    conn.close()
    print(f"Dados para {dados['nome_comum']} inseridos no banco de dados 'paises.db'.")

# --- PARTE 2: Web Scraping com BeautifulSoup (Livros) ---

def configurar_banco_dados_livros():
    conn = sqlite3.connect('livraria.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS livros (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            titulo TEXT,
            preco REAL,
            avaliacao_estrelas TEXT,
            disponibilidade TEXT
        )
    ''')
    conn.commit()
    conn.close()
    print("Banco de dados 'livraria.db' e tabela 'livros' verificados/criados.")

def raspar_dados_livros():
    url = "https://books.toscrape.com/"
    print(f"Acessando o site para raspar livros: {url}")
    try:
        response = requests.get(url)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(f"Erro ao acessar a página de livros: {e}")
        return []

    soup = BeautifulSoup(response.text, 'html.parser')
    
    livros_encontrados = []
    for livro_tag in soup.find_all('article', class_='product_pod')[:10]:
        titulo = livro_tag.h3.a['title']
        
        preco_texto = livro_tag.find('p', class_='price_color').get_text()
        preco = float(''.join(filter(lambda x: x.isdigit() or x == '.', preco_texto)))

        avaliacao_tag = livro_tag.find('p', class_='star-rating')
        avaliacao_estrelas = avaliacao_tag['class'][1] if avaliacao_tag and len(avaliacao_tag['class']) > 1 else 'N/A'
        
        disponibilidade_tag = livro_tag.find('p', class_='instock availability')
        disponibilidade = disponibilidade_tag.get_text(strip=True) if disponibilidade_tag else 'N/A'
        
        livros_encontrados.append({
            "titulo": titulo,
            "preco": preco,
            "avaliacao_estrelas": avaliacao_estrelas,
            "disponibilidade": disponibilidade
        })
    
    return livros_encontrados

def inserir_dados_livros(livros_data):
    conn = sqlite3.connect('livraria.db')
    cursor = conn.cursor()
    for livro in livros_data:
        cursor.execute('''
            INSERT INTO livros (titulo, preco, avaliacao_estrelas, disponibilidade)
            VALUES (?, ?, ?, ?)
        ''', (
            livro['titulo'],
            livro['preco'],
            livro['avaliacao_estrelas'],
            livro['disponibilidade']
        ))
    conn.commit()
    conn.close()
    print(f"{len(livros_data)} livros inseridos no banco de dados 'livraria.db'.")

# --- PARTE 3: Relatório Final (Word) ---

def gerar_relatorio_word_com_input_nome():
    document = Document()

    nome_aluno = input("Por favor, digite seu nome completo para o relatório: ").strip()
    if not nome_aluno:
        nome_aluno = "Nome do Aluno Desconhecido"

    document.add_heading('Relatório de Dados Públicos', level=1)
    
    data_geracao = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    document.add_paragraph(f"Nome do Aluno: {nome_aluno}")
    document.add_paragraph(f"Data de Geração: {data_geracao}")

    document.add_heading('1. Dados dos Países', level=2)
    conn_paises = sqlite3.connect('paises.db')
    cursor_paises = conn_paises.cursor()
    cursor_paises.execute("SELECT * FROM paises")
    paises_data = cursor_paises.fetchall()
    conn_paises.close()

    if paises_data:
        colunas_paises = [
            "ID", "Nome Comum", "Nome Oficial", "Capital", "Continente", 
            "Região", "Sub-região", "População", "Área (km²)", "Moeda", 
            "Símbolo Moeda", "Idioma Principal", "Fuso Horário", "URL Bandeira"
        ]
        
        table_paises = document.add_table(rows=1, cols=len(colunas_paises))
        table_paises.style = 'Table Grid'
        hdr_cells_paises = table_paises.rows[0].cells
        for i, header_text in enumerate(colunas_paises):
            hdr_cells_paises[i].text = header_text

        for row_data in paises_data:
            row_cells = table_paises.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
    else:
        document.add_paragraph("Nenhum dado de país encontrado para incluir no relatório.")

    document.add_page_break()

    document.add_heading('2. Dados dos Livros', level=2)
    conn_livros = sqlite3.connect('livraria.db')
    cursor_livros = conn_livros.cursor()
    cursor_livros.execute("SELECT * FROM livros")
    livros_data = cursor_livros.fetchall()
    conn_livros.close()

    if livros_data:
        colunas_livros = ["ID", "Título", "Preço", "Avaliação", "Disponibilidade"]
        
        table_livros = document.add_table(rows=1, cols=len(colunas_livros))
        table_livros.style = 'Table Grid'
        hdr_cells_livros = table_livros.rows[0].cells
        for i, header_text in enumerate(colunas_livros):
            hdr_cells_livros[i].text = header_text

        for row_data in livros_data:
            row_cells = table_livros.add_row().cells
            for i, cell_data in enumerate(row_data):
                if colunas_livros[i] == "Preço":
                    row_cells[i].text = f"£{cell_data:.2f}"
                else:
                    row_cells[i].text = str(cell_data)
    else:
        document.add_paragraph("Nenhum dado de livro encontrado para incluir no relatório.")

    nome_arquivo = 'relatorio_final.docx'
    document.save(nome_arquivo)
    print(f"\nRelatório '{nome_arquivo}' gerado com sucesso!")

# --- Função Principal para Orquestrar as Partes ---

def main_orchestrator():
    print("--- Iniciando Parte 1: Extração de Dados de Países via API REST ---")
    configurar_banco_dados_paises()
    
    paises_para_buscar = []
    for i in range(3):
        nome_pais = input(f"Digite o nome do país {i+1} (em inglês): ").strip()
        paises_para_buscar.append(nome_pais)

    for nome_pais in paises_para_buscar:
        print(f"\nBuscando dados para {nome_pais}...")
        dados_pais = obter_dados_pais(nome_pais)
        if dados_pais:
            inserir_dados_pais(dados_pais)
        else:
            print(f"Não foi possível obter dados para {nome_pais}. Pulando a inserção no banco de dados.")

    print("\n--- Parte 1 Concluída. ---")

    print("\n--- Iniciando Parte 2: Web Scraping de Livros com BeautifulSoup ---")
    configurar_banco_dados_livros()
    
    dados_livros = raspar_dados_livros()
    
    if dados_livros:
        inserir_dados_livros(dados_livros)
    else:
        print("Nenhum dado de livro foi coletado para inserção.")

    print("\n--- Parte 2 Concluída. ---")
    print("\n--- Iniciando Parte 3: Geração do Relatório Final ---")
    gerar_relatorio_word_com_input_nome()
    print("\n--- Parte 3 Concluída. ---")
    print("\nProcesso completo finalizado. Verifique 'paises.db', 'livraria.db' e 'relatorio_final.docx'.")

if __name__ == "__main__":
    main_orchestrator()